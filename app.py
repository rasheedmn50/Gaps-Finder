#!/usr/bin/env python3
# Streamlit Paper Gap Analyzer — clean 2-paragraph DOCX + OCR + clustering + CSV + appendix
# -----------------------------------------------------------------------------------------
# Upload up to 50 PDFs/DOCX/TXT. Analyze shared research gaps, cluster similar gaps,
# and export a clean, branded **two-paragraph** Word report (plus other styles).

from __future__ import annotations

import io
import json
import re
import time
import warnings
from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st

# Optional deps
try:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadWarning
    warnings.filterwarnings("ignore", category=PdfReadWarning)
except Exception:
    PdfReader = None

try:
    import pytesseract
    from PIL import Image  # noqa: F401
    import pypdfium2 as pdfium
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

try:
    from sklearn.preprocessing import normalize as sk_normalize
except Exception:
    sk_normalize = None

try:
    import hdbscan
    HDBSCAN_AVAILABLE = True
except Exception:
    HDBSCAN_AVAILABLE = False

from openai import OpenAI

# ------------------------------------------------------------
# UI + Defaults
# ------------------------------------------------------------
st.set_page_config(page_title="Paper Gap Analyzer", page_icon="🧠", layout="wide")

DEFAULT_MODEL = "gpt-5"
DEFAULT_EMBEDDING = "text-embedding-3-large"
MAX_PAPERS = 50
MAX_CHARS_PER_PAPER = 80_000

def _safe_rerun():
    try:
        st.rerun()
    except Exception:
        try:
            st.experimental_rerun()
        except Exception:
            pass

ss = st.session_state
if "stage" not in ss:
    ss.stage = "upload"
    ss.files = []
    ss.per_paper = []
    ss.accepted = []
    ss.stats = []
    ss.order = []
    ss.payload = {}
    ss.appendix = []

# ------------------------------------------------------------
# Sidebar
# ------------------------------------------------------------
st.sidebar.title("⚙️ Settings")

api_key_input = st.sidebar.text_input("OpenAI API Key", type="password", placeholder="sk-…")
api_key = api_key_input or st.secrets.get("OPENAI_API_KEY", "")

model = st.sidebar.text_input("Chat Model", value=DEFAULT_MODEL)
embedding_model = st.sidebar.text_input("Embedding Model", value=DEFAULT_EMBEDDING)
limit_chars = st.sidebar.number_input("Max chars per paper", 20_000, 180_000, MAX_CHARS_PER_PAPER, 2_000)

st.sidebar.markdown("---")
rigor = st.sidebar.toggle("High-Rigor mode", value=True,
                          help="Evidence quotes + page markers + section-aware prompts + HDBSCAN if available.")
sim_threshold = st.sidebar.slider("Similarity threshold (greedy clustering)", 0.70, 0.92, 0.82, 0.01)

st.sidebar.markdown("---")
tagging_on = st.sidebar.toggle("Checklist tagging (extra LLM calls)", value=False)

st.sidebar.markdown("---")
st.sidebar.subheader("🖼 Branding")
logo_upload = st.sidebar.file_uploader("Logo (PNG/JPG)", type=["png", "jpg", "jpeg"])
accent = st.sidebar.text_input("Accent hex color", value="#3B82F6")

st.sidebar.markdown("---")
st.sidebar.subheader("🔎 OCR for scanned PDFs")
ocr_mode = st.sidebar.selectbox(
    "OCR mode",
    options=["Auto (OCR blank/low-text pages)", "Force OCR (all PDF pages)", "Off"],
    index=0 if OCR_AVAILABLE else 2,
    help="OCR uses pypdfium2 + pytesseract."
)

st.sidebar.markdown("✅ DOCX: **available**" if Document else "⚠️ DOCX: **unavailable** (RTF/TXT fallback)")

# ------------------------------------------------------------
# Branding helpers
# ------------------------------------------------------------
DEFAULT_LOGO_PATH = "assets/logo.png"

def get_logo_bytes() -> Optional[bytes]:
    if logo_upload is not None:
        try:
            return logo_upload.getvalue()
        except Exception:
            pass
    try:
        with open(DEFAULT_LOGO_PATH, "rb") as f:
            return f.read()
    except Exception:
        return None

def inject_brand_css(accent_hex: str):
    # Plain string (not f-string) to avoid curly-brace issues
    css = (
        "<style>\n"
        ".app-title {\n"
        "  background: linear-gradient(90deg, " + accent_hex + "22, transparent);\n"
        "  border-left: 6px solid " + accent_hex + ";\n"
        "  padding: 10px 14px;\n"
        "  border-radius: 10px;\n"
        "  margin-bottom: 10px;\n"
        "}\n"
        "</style>"
    )
    st.markdown(css, unsafe_allow_html=True)

inject_brand_css(accent)
logo_bytes = get_logo_bytes()

# Header
left, right = st.columns([1, 5])
with left:
    if logo_bytes:
        st.image(logo_bytes, width="stretch")
with right:
    st.markdown('<div class="app-title"><h1 style="margin:0;">🧠 Paper Gap Analyzer</h1></div>', unsafe_allow_html=True)
    st.caption("Find shared research gaps across uploaded papers and export a clean, branded two-paragraph Word report.")

# ------------------------------------------------------------
# Parsers & LLM helpers
# ------------------------------------------------------------
def _to_float(s):
    try:
        return float(s)
    except Exception:
        return None

def parse_confidence(x) -> float:
    if isinstance(x, (int, float)):
        val = float(x)
        if val > 1.0:
            val = val / 100.0 if val <= 100 else 1.0
        return max(0.0, min(1.0, val))
    if isinstance(x, str):
        s = x.strip().lower()
        m = re.match(r'^(\d+(?:\.\d+)?)\s*%$', s)
        if m:
            return max(0.0, min(1.0, float(m.group(1)) / 100.0))
        n = _to_float(s)
        if n is not None:
            if n > 1.0:
                n = n / 100.0 if n <= 100 else 1.0
            return max(0.0, min(1.0, n))
        mapping = {"very low": 0.15, "low": 0.25, "medium": 0.5, "med": 0.5, "moderate": 0.5,
                   "high": 0.85, "very high": 0.95}
        for k, v in mapping.items():
            if k in s:
                return v
    return 0.5

def parse_severity(x) -> int:
    if isinstance(x, (int, float)):
        return int(max(1, min(5, round(float(x)))))
    if isinstance(x, str):
        s = x.strip().lower()
        n = _to_float(s)
        if n is not None:
            return int(max(1, min(5, round(n))))
        mapping = {"trivial": 1, "minor": 2, "low": 2, "medium": 3, "med": 3, "moderate": 3,
                   "high": 4, "major": 4, "very high": 5, "critical": 5, "severe": 5}
        for k, v in mapping.items():
            if k in s:
                return v
    return 3

def build_client(key: str) -> OpenAI:
    if not key or not key.startswith("sk-"):
        raise RuntimeError("Please provide a valid OpenAI API key.")
    return OpenAI(api_key=key)

def cleanup_json_str(s: str) -> str:
    s = s.strip()
    s = re.sub(r"^```(json)?\s*|```$", "", s, flags=re.IGNORECASE | re.DOTALL).strip()
    s = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", " ", s)
    m = re.search(r"\{[\s\S]*\}", s)
    if m:
        s = m.group(0)
    s = re.sub(r",\s*(\])", r"\1", s)
    s = re.sub(r",\s*(\})", r"\1", s)
    return s

def parse_possible_json(content: str) -> Dict[str, Any]:
    try:
        return json.loads(content)
    except Exception:
        return json.loads(cleanup_json_str(content))

def safe_chat(client: OpenAI, model: str, messages: List[Dict[str, str]], *,
              use_schema: bool, schema: Optional[Dict[str, Any]] = None, retries: int = 4) -> str:
    last_err = None
    allow_plain = False
    for attempt in range(1, retries + 1):
        try:
            msgs = [dict(m) for m in messages]
            if use_schema and not allow_plain and schema:
                resp = client.chat.completions.create(
                    model=model, messages=msgs,
                    response_format={"type": "json_schema", "json_schema": schema},
                )
            else:
                if msgs and "Return ONLY valid JSON" not in msgs[-1]["content"]:
                    msgs[-1]["content"] += "\n\nReturn ONLY valid, minified JSON that matches the expected fields."
                resp = client.chat.completions.create(model=model, messages=msgs)
            return resp.choices[0].message.content
        except Exception as e:
            err = str(e).lower()
            last_err = e
            if "unsupported" in err or "response_format" in err or "temperature" in err:
                allow_plain = True
            time.sleep(1.0 * attempt)
    raise RuntimeError(f"LLM request failed after retries: {last_err}")

# Schemas & normalization
def paper_schema_simple(name="paper_gap") -> Dict[str, Any]:
    return {
        "name": name,
        "strict": True,
        "schema": {
            "type": "object", "additionalProperties": True,
            "properties": {
                "paper_id": {"type": "string"},
                "title": {"type": "string"},
                "domain": {"type": "string"},
                "topic_keywords": {"type": "array", "items": {"type": "string"}},
                "methodology": {"type": "string"},
                "reported_limitations": {"type": "array", "items": {"type": ["string", "object"]}},
                "reported_future_work": {"type": "array", "items": {"type": ["string", "object"]}},
                "inferred_gaps": {"type": "array", "items": {"type": ["string", "object"]}}
            },
            "required": ["paper_id", "title", "domain", "topic_keywords", "methodology",
                         "reported_limitations", "reported_future_work", "inferred_gaps"]
        }
    }

def paper_schema_rigorous(name="paper_gap_v2") -> Dict[str, Any]:
    return {
        "name": name,
        "strict": True,
        "schema": {
            "type": "object",
            "additionalProperties": True,
            "properties": {
                "paper_id": {"type": "string"},
                "title": {"type": "string"},
                "domain": {"type": "string"},
                "topic_keywords": {"type": "array", "items": {"type": "string"}},
                "methodology": {"type": "string"},
                "reported_limitations": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": True,
                        "properties": {
                            "description": {"type": "string"},
                            "section": {"type": "string"},
                            "evidence_quote": {"type": "string"},
                            "evidence_page": {"type": ["integer", "array", "string", "null"]},
                            "confidence": {"type": ["number", "integer"]}
                        },
                        "required": ["description", "section", "evidence_quote"]
                    }
                },
                "reported_future_work": {"type": "array", "items": {"type": ["string", "object"]}},
                "inferred_gaps": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": True,
                        "properties": {
                            "description": {"type": "string"},
                            "category": {"type": "string"},
                            "severity": {"type": ["integer", "number"]},
                            "confidence": {"type": ["number", "integer"]},
                            "evidence_quote": {"type": ["string", "null"]},
                            "evidence_page": {"type": ["integer", "array", "string", "null"]}
                        },
                        "required": ["description", "category", "severity", "confidence"]
                    }
                }
            },
            "required": [
                "paper_id", "title", "domain", "topic_keywords", "methodology",
                "reported_limitations", "reported_future_work", "inferred_gaps"
            ]
        }
    }

def norm_list_str(x):
    if not isinstance(x, list):
        return []
    return [str(t) for t in x if isinstance(t, (str, int, float))]

def normalize_simple(data: Dict[str, Any], paper_id: str) -> Dict[str, Any]:
    out = {
        "paper_id": str(data.get("paper_id", paper_id)),
        "title": str(data.get("title", paper_id)),
        "domain": str(data.get("domain", "")),
        "topic_keywords": norm_list_str(data.get("topic_keywords") or []),
        "methodology": str(data.get("methodology", "")),
        "reported_limitations": [], "reported_future_work": [], "inferred_gaps": []
    }
    for k in ("reported_limitations", "reported_future_work"):
        v = data.get(k) or []
        cleaned = []
        if isinstance(v, list):
            for item in v:
                if isinstance(item, str):
                    s = item.strip()
                    if s:
                        cleaned.append(s)
                elif isinstance(item, dict):
                    s = str(item.get("text") or item.get("description") or "").strip()
                    if s:
                        cleaned.append(s)
        out[k] = cleaned
    gaps = data.get("inferred_gaps") or []
    if isinstance(gaps, list):
        for g in gaps:
            if isinstance(g, dict):
                desc = str(g.get("description") or g.get("text") or "").strip()
                if not desc:
                    continue
                sev = parse_severity(g.get("severity", 3))
                conf = parse_confidence(g.get("confidence", 0.5))
                out["inferred_gaps"].append({
                    "description": desc, "category": str(g.get("category") or "Other"),
                    "severity": min(max(sev, 1), 5), "confidence": conf,
                    "evidence_quote": g.get("evidence_quote") if isinstance(g.get("evidence_quote"), str) else None,
                    "evidence_page": g.get("evidence_page")
                })
            elif isinstance(g, (str, int, float)):
                s = str(g).strip()
                if s:
                    out["inferred_gaps"].append({
                        "description": s, "category": "Other",
                        "severity": 3, "confidence": 0.5,
                        "evidence_quote": None, "evidence_page": None
                    })
    return out

def normalize_rigorous(data: Dict[str, Any], paper_id: str) -> Dict[str, Any]:
    out = normalize_simple(data, paper_id)
    rep = []
    for item in data.get("reported_limitations") or []:
        if isinstance(item, dict):
            desc = str(item.get("description") or item.get("text") or "").strip()
            if not desc:
                continue
            rep.append({
                "description": desc, "section": str(item.get("section") or ""),
                "evidence_quote": str(item.get("evidence_quote") or "")[:400],
                "evidence_page": item.get("evidence_page"),
                "confidence": parse_confidence(item.get("confidence", 0.8))
            })
        elif isinstance(item, (str, int, float)):
            rep.append({
                "description": str(item).strip(), "section": "",
                "evidence_quote": "", "evidence_page": None, "confidence": 0.6
            })
    out["reported_limitations"] = rep
    return out

# ------------------------------------------------------------
# File reading + OCR
# ------------------------------------------------------------
def pdf_text_pages_pypdf(data: bytes) -> List[str]:
    if PdfReader is None:
        raise RuntimeError("pypdf not installed")
    try:
        try:
            reader = PdfReader(io.BytesIO(data), strict=False)
        except TypeError:
            reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return pages
    except Exception as e:
        raise RuntimeError(f"PDF parse failed: {e}")

def ocr_pages_pdfium(data: bytes, dpi: int = 220, lang: str = "eng") -> List[str]:
    if not OCR_AVAILABLE:
        return []
    try:
        pdf = pdfium.PdfDocument(io.BytesIO(data))
        n = len(pdf)
        results = []
        for i in range(n):
            page = pdf[i]
            bitmap = page.render(scale=dpi / 72.0)
            pil_image = bitmap.to_pil()
            txt = pytesseract.image_to_string(pil_image, lang=lang)
            results.append(txt or "")
        return results
    except Exception:
        return []

def read_docx_stream(file_bytes: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    d = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in d.paragraphs)

def extract_text_with_pages_from_bytes(name: str, data: bytes, ocr_mode: str) -> Tuple[str, List[str]]:
    lname = name.lower()
    if lname.endswith(".txt"):
        t = data.decode("utf-8", errors="ignore")
        return t, [t]
    if lname.endswith(".docx"):
        t = read_docx_stream(data)
        return t, [t]
    if lname.endswith(".pdf"):
        pypdf_pages = pdf_text_pages_pypdf(data)
        if ocr_mode == "Off" or not OCR_AVAILABLE:
            pages = pypdf_pages
        elif ocr_mode == "Force OCR (all PDF pages)":
            pages = ocr_pages_pdfium(data) or pypdf_pages
        else:
            pages = pypdf_pages[:]
            to_ocr = [i for i, t in enumerate(pages) if len((t or "").strip()) < 40]
            if to_ocr:
                ocr_all = ocr_pages_pdfium(data)
                if ocr_all:
                    for i in to_ocr:
                        if i < len(ocr_all) and len((ocr_all[i] or "").strip()) > len((pages[i] or "").strip()):
                            pages[i] = ocr_all[i]
        marked = [f"[[PAGE {i}]]\n{p}" for i, p in enumerate(pages, 1)]
        return "\n\n".join(marked), pages
    raise ValueError("Unsupported file type")

# ------------------------------------------------------------
# Sectioning (heuristic)
# ------------------------------------------------------------
SECTION_HEADERS = [
    r"\babstract\b", r"\bbackground\b", r"\bintroduction\b", r"\bmethods?\b",
    r"\bmaterials and methods\b", r"\bresults\b", r"\bdiscussion\b",
    r"\blimitations?\b", r"\bconclusion(s)?\b"
]

def split_into_sections(text: str) -> Dict[str, str]:
    idxs = []
    for m in re.finditer("|".join(SECTION_HEADERS), text, flags=re.IGNORECASE):
        idxs.append((m.start(), m.group(0).lower()))
    if not idxs:
        return {"FullText": text}
    idxs.sort()
    sections: Dict[str, str] = {}
    for i, (pos, hdr) in enumerate(idxs):
        end = idxs[i + 1][0] if i + 1 < len(idxs) else len(text)
        sections[hdr.strip().title()] = text[pos:end].strip()
    return sections

# ------------------------------------------------------------
# LLM Analysis
# ------------------------------------------------------------
def analyze_paper_simple(client: OpenAI, model: str, paper_id: str, text: str) -> Dict[str, Any]:
    if len(text) > MAX_CHARS_PER_PAPER:
        text = text[:MAX_CHARS_PER_PAPER]
    system = ("You are a meticulous research-methods auditor.\n"
              "Extract concise structured metadata, reported limitations, and inferred gaps grounded in the paper text.")
    user = ("IMPORTANT: `severity` = integer 1–5 (no words). `confidence` = number 0–1 (no words; default 0.5).\n\n"
            f"PAPER_ID: {paper_id}\n\nReturn only JSON matching the schema (no prose).\n\n"
            f"--- PAPER FULL TEXT START ---\n{text}\n--- PAPER FULL TEXT END ---")
    content = safe_chat(client, model,
                        [{"role": "system", "content": system}, {"role": "user", "content": user}],
                        use_schema=True, schema=paper_schema_simple())
    return normalize_simple(parse_possible_json(content), paper_id)

def analyze_paper_rigorous(client: OpenAI, model: str, paper_id: str, sectioned_text: Dict[str, str]) -> Dict[str, Any]:
    parts = []
    for k, v in sectioned_text.items():
        if not v or not v.strip():
            continue
        vv = v.strip()
        budget = max(2000, int(MAX_CHARS_PER_PAPER / max(1, len(sectioned_text))))
        parts.append(f"=== {k} ===\n{vv[:budget]}\n")
    joined = "\n".join(parts)

    system = ("You are an evidence-first research auditor. Use quotes and page markers when available.\n"
              "If [[PAGE N]] markers appear near quoted text, copy N into 'evidence_page'. "
              "For inferred gaps, keep confidence low unless supported by a quote.")
    user = ("IMPORTANT: `severity` = integer 1–5. `confidence` = number 0–1 (default 0.5).\n\n"
            "Return ONLY JSON matching the schema. For each REPORTED limitation: include 'description', 'section', "
            "'evidence_quote' (<=60 words), and 'evidence_page' from any [[PAGE N]] markers if possible.\n"
            "For each INFERRED gap: include 'description','category','severity (1-5)','confidence', and optional 'evidence_quote'/'evidence_page'.\n\n"
            + joined)
    content = safe_chat(client, model,
                        [{"role": "system", "content": system}, {"role": "user", "content": user}],
                        use_schema=True, schema=paper_schema_rigorous())
    return normalize_rigorous(parse_possible_json(content), paper_id)

# ------------------------------------------------------------
# Embeddings & clustering
# ------------------------------------------------------------
def embed(client: OpenAI, embedding_model: str, texts: List[str]) -> np.ndarray:
    out = []
    if not texts:
        return np.zeros((0, 1536), dtype=np.float32)
    B = 64
    for i in range(0, len(texts), B):
        batch = texts[i:i + B]
        emb = client.embeddings.create(model=embedding_model, input=batch)
        vecs = [d.embedding for d in emb.data]
        out.append(np.array(vecs, dtype=np.float32))
    return np.vstack(out)

def greedy_cluster(vectors: np.ndarray, threshold: float) -> List[List[int]]:
    n = len(vectors)
    unused = set(range(n))
    clusters: List[List[int]] = []
    while unused:
        seed = unused.pop()
        members = [seed]
        centroid = vectors[seed].copy()
        changed = True
        while changed:
            changed = False
            to_add = []
            for j in list(unused):
                sim = float(np.dot(centroid, vectors[j]) / (np.linalg.norm(centroid) * np.linalg.norm(vectors[j]) + 1e-9))
                if sim >= threshold:
                    to_add.append(j)
            if to_add:
                for j in to_add:
                    unused.remove(j)
                    members.append(j)
                centroid = vectors[members].mean(axis=0)
                changed = True
        clusters.append(members)
    return clusters

def hdbscan_cluster(vectors: np.ndarray) -> List[List[int]]:
    if not HDBSCAN_AVAILABLE or sk_normalize is None:
        return greedy_cluster(vectors, threshold=0.82)
    X = sk_normalize(vectors)
    clusterer = hdbscan.HDBSCAN(min_cluster_size=3, min_samples=1, metric='euclidean')
    labels = clusterer.fit_predict(X)
    clusters_map: Dict[int, List[int]] = {}
    for idx, lab in enumerate(labels):
        if lab == -1:
            clusters_map[f"noise-{idx}"] = [idx]
        else:
            clusters_map.setdefault(lab, []).append(idx)
    return list(clusters_map.values())

def cluster_stats_from_members(members: List[int], items: List[Dict[str, Any]]) -> Dict[str, Any]:
    subset = [items[i] for i in members]
    uniq_papers = sorted({x["paper_id"] for x in subset})
    severities = [parse_severity(x.get("severity", 3)) for x in subset]
    srcs = [x.get("source", "Reported") for x in subset]
    evidence = [1 if (x.get("evidence_quote")) else 0 for x in subset]
    return {
        "paper_coverage": len(uniq_papers),
        "unique_papers": uniq_papers,
        "count": len(subset),
        "avg_severity": float(np.mean(severities) if severities else 0),
        "median_severity": float(np.median(severities) if severities else 0),
        "reported_share": float(srcs.count("Reported") / len(srcs)) if srcs else 0.0,
        "inferred_share": float(srcs.count("Inferred") / len(srcs)) if srcs else 0.0,
        "evidence_strength": float(np.mean(evidence) if evidence else 0.0),
        "samples": [x["text"] for x in subset[:5]],
        "members_idx": members
    }

def rank_clusters(stats: List[Dict[str, Any]], total_papers: int, rigorous: bool) -> List[int]:
    scores = []
    for i, c in enumerate(stats):
        coverage = c["paper_coverage"] / max(1, total_papers)
        sev = c["median_severity"] / 5.0
        ev = c.get("evidence_strength", 0.0)
        score = (0.5 * coverage + 0.3 * sev + 0.2 * ev) if rigorous else (0.7 * coverage + 0.3 * sev)
        scores.append((score, i))
    scores.sort(reverse=True)
    return [i for _, i in scores]

# ------------------------------------------------------------
# Optional: checklist tagging LLM
# ------------------------------------------------------------
CHECKLIST = [
    "ExternalValidation", "ProspectiveValidation", "Calibration", "Fairness/Equity",
    "Reproducibility", "ReportingCompleteness", "Safety/ClinicalRisk", "DataShift/Generalizability",
    "ComparatorAdequacy", "Confounding/Bias", "Ethics/Governance", "Privacy"
]

def llm_tag_gaps(client: OpenAI, model: str, gap_texts: List[str]) -> List[List[str]]:
    if not gap_texts:
        return [[] for _ in gap_texts]
    system = "You assign concise checklist tags to research gaps. Return ONLY JSON."
    user = ("For each gap, choose zero or more tags from this list:\n" + ", ".join(CHECKLIST) +
            "\nReturn JSON with {'tags': [['Tag1','Tag2',...], ...]} in the same order as inputs.\n\n" +
            json.dumps({"gaps": gap_texts}, ensure_ascii=False))
    content = safe_chat(
        client, model,
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        use_schema=False
    )
    try:
        data = parse_possible_json(content)
        return data.get("tags", [[] for _ in gap_texts])
    except Exception:
        return [[] for _ in gap_texts]

# ------------------------------------------------------------
# Two-paragraph writer + sanitizer
# ------------------------------------------------------------
def _strip_code_fences(s: str) -> str:
    return re.sub(r"^```+.*?```+$", "", s.strip(), flags=re.DOTALL)

def _remove_json_wrappers(s: str) -> str:
    s2 = s.strip()
    m = re.match(r'^\s*\{\s*"text"\s*:\s*"(.*)"\s*\}\s*$', s2, flags=re.DOTALL)
    if m:
        return m.group(1)
    m = re.match(r'^\s*\{\s*"(?:paragraphs?|paras?)"\s*:\s*\[(.*)\]\s*\}\s*$', s2, flags=re.DOTALL)
    if m:
        inner = m.group(1)
        parts = re.findall(r'"([^"]+)"', inner)
        return "\n\n".join(parts)
    return s

def _basic_deescape(s: str) -> str:
    s = s.replace("\\n", "\n")
    if s.startswith('"') and s.endswith('"'):
        s = s[1:-1]
    return s

def clean_to_two_paragraphs(text: str) -> List[str]:
    s = _strip_code_fences(text)
    s = _remove_json_wrappers(s)
    s = _basic_deescape(s)
    s = re.sub(r'^\s*\{|\}\s*$', '', s.strip())
    s = re.sub(r"\n{3,}", "\n\n", s)
    paras = [p.strip() for p in re.split(r"\n\s*\n", s) if p.strip()]
    if len(paras) >= 2:
        return [paras[0], paras[1]]
    if len(paras) == 1:
        sentences = re.split(r'(?<=[\.\!\?])\s+', paras[0])
        mid = max(1, len(sentences)//2)
        p1 = " ".join(sentences[:mid]).strip()
        p2 = " ".join(sentences[mid:]).strip()
        return [p1, p2] if p2 else [p1, "This paragraph explains how to build on the most important gap with a clear plan."]
    return ["", ""]

def llm_two_paragraph_report(client: OpenAI, model: str, clusters_payload: Dict[str, Any]) -> List[str]:
    system = ("You write clear, pleasant prose in very simple words for busy researchers. "
              "Return plain text only. No lists, no headings, no bullets, no code, no JSON, no braces.")
    user = ("Write EXACTLY two short paragraphs.\n"
            "Paragraph 1: explain the main gaps shared across these papers in simple words and CLEARLY name the one most important gap.\n"
            "Paragraph 2: explain, in simple steps, how to build on that important gap to create genuine novelty for a future paper.\n"
            "Rules: no lists, no bullets, no headings, no numbers or percentages, no quotes, no curly braces. Plain sentences only.\n\n"
            "JSON INPUT YOU SHOULD READ (do not copy):\n" + json.dumps(clusters_payload, ensure_ascii=False))
    content = safe_chat(
        client, model,
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        use_schema=False)
    return clean_to_two_paragraphs(content)

# Other styles
def llm_narrative(client: OpenAI, model: str, clusters_payload: Dict[str, Any]) -> str:
    system = "You write clear, non-technical academic summaries. Return plain text only."
    user = ("Write 3–5 short paragraphs narrating the main shared gaps, "
            "then ONE final paragraph recommending the single gap with the best potential, with a brief rationale and a one-sentence next step.\n"
            "Rules: NO lists, NO bullets, NO headings, NO explicit numbers/percentages—just prose. <= 600 words.\n\n"
            "JSON:\n" + json.dumps(clusters_payload, ensure_ascii=False))
    content = safe_chat(client, model,
                        [{"role": "system", "content": system}, {"role": "user", "content": user}],
                        use_schema=False)
    return re.sub(r"`{3,}.*?`{3,}", "", content, flags=re.DOTALL).strip()

def llm_executive_brief(client: OpenAI, model: str, clusters_payload: Dict[str, Any]) -> str:
    system = "You write ultra-concise executive briefs for researchers."
    user = ("In <= 220 words total, write two short paragraphs that plainly describe the main gaps, "
            "then one final 2–3 sentence recommendation naming the best gap and a concrete starting step.\n"
            "No headings, no bullets.\n\n"
            "JSON:\n" + json.dumps(clusters_payload, ensure_ascii=False))
    content = safe_chat(client, model,
                        [{"role": "system", "content": system}, {"role": "user", "content": user}],
                        use_schema=False)
    return re.sub(r"`{3,}.*?`{3,}", "", content, flags=re.DOTALL).strip()

def llm_plain_language(client: OpenAI, model: str, clusters_payload: Dict[str, Any]) -> str:
    system = "You explain research in very simple words. Short sentences. No jargon. Plain text only."
    user = ("Write 3–4 short paragraphs in very simple words explaining the main shared gaps. "
            "Then one final paragraph naming the best gap to work on next and a simple way to start. "
            "No lists, bullets, headings, or numbers. <= 400 words.\n\n"
            "JSON:\n" + json.dumps(clusters_payload, ensure_ascii=False))
    content = safe_chat(client, model,
                        [{"role": "system", "content": system}, {"role": "user", "content": user}],
                        use_schema=False)
    return re.sub(r"`{3,}.*?`{3,}", "", content, flags=re.DOTALL).strip()

# Writers
def docx_from_paragraphs_with_appendix(title, paragraphs, appendix_items, logo_bytes, accent_hex) -> bytes:
    if Document is None:
        return b""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    if logo_bytes:
        try:
            doc.add_picture(io.BytesIO(logo_bytes), width=Inches(1.4))
        except Exception:
            pass
    h = doc.add_paragraph(title); h.style = "Title"; h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"); sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    for ptxt in [p.strip() for p in paragraphs if p.strip()]:
        p = doc.add_paragraph(ptxt)
        pf = p.paragraph_format; pf.space_after = Pt(8); pf.line_spacing = 1.25
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if appendix_items:
        doc.add_paragraph()
        app_h = doc.add_paragraph("Appendix — Representative Quotes")
        if app_h.runs: app_h.runs[0].bold = True
        for it in appendix_items[:40]:
            line = f"{it.get('paper_id', 'unknown')} — "
            if it.get("evidence_page"): line += f"[p. {it.get('evidence_page')}] "
            line += f"“{(it.get('evidence_quote') or '').strip()}”"
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(4); para.paragraph_format.line_spacing = 1.15
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def rtf_from_paragraphs(title: str, paras: List[str]) -> bytes:
    def esc(s: str) -> str:
        return s.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")
    body = "".join([esc(p) + r"\par " for p in paras])
    return (r"{\rtf1\ansi\deff0\fs28 \b " + esc(title) + r"\b0\par " + body + "}").encode("utf-8")

def txt_from_paragraphs(title: str, paras: List[str]) -> bytes:
    return ("\n\n".join([title, ""] + paras)).encode("utf-8")

# ------------------------------------------------------------
# Upload + Controls
# ------------------------------------------------------------
uploaded = st.file_uploader("Upload your papers (PDF / DOCX / TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
col_u1, col_u2 = st.columns([1, 1])
start_btn = col_u1.button("Analyze", type="primary", width="stretch")
reset_btn = col_u2.button("Reset", width="stretch")

if reset_btn:
    ss.clear()
    _safe_rerun()

if uploaded:
    files_list = []
    for uf in uploaded[:MAX_PAPERS]:
        try:
            data = uf.read()
            files_list.append({"name": uf.name, "bytes": data, "type": uf.type})
        except Exception:
            pass
    if files_list:
        ss.files = files_list

# ------------------------------------------------------------
# Step 1 — Analyze papers
# ------------------------------------------------------------
if start_btn:
    if not api_key:
        st.error("Please enter your OpenAI API key.")
        st.stop()
    if not ss.files:
        st.warning("Please upload at least one file.")
        st.stop()

    client = build_client(api_key)
    per_paper: List[Dict[str, Any]] = []
    progress = st.progress(0.0)
    status = st.empty()

    for i, f in enumerate(ss.files, 1):
        status.info(f"Reading: {f['name']} (OCR: {ocr_mode})")
        try:
            full_text, _pages = extract_text_with_pages_from_bytes(f["name"], f["bytes"], ocr_mode)
        except Exception as e:
            st.warning(f"Skipping {f['name']}: read error — {e}")
            progress.progress(i / len(ss.files)); continue

        if len(full_text.strip()) < 500:
            st.warning(f"Skipping {f['name']}: very little text (try OCR Force / different language).")
            progress.progress(i / len(ss.files)); continue

        status.info(f"LLM analysis: {f['name']} ({'rigorous' if rigor else 'simple'})")
        try:
            if rigor:
                sections = split_into_sections(full_text)
                analysis = analyze_paper_rigorous(client, model, f["name"], sections)
            else:
                analysis = analyze_paper_simple(client, model, f["name"], full_text[:limit_chars])
        except Exception as e:
            st.warning(f"LLM analysis failed for {f['name']}: {e}")
            progress.progress(i / len(ss.files)); continue

        per_paper.append(analysis)
        progress.progress(i / len(ss.files))

    status.empty()
    if not per_paper:
        st.error("No papers analyzed successfully.")
        st.stop()

    ss.per_paper = per_paper
    ss.stage = "review"
    _safe_rerun()

# ------------------------------------------------------------
# Step 2 — Review (ONLY when stage == 'review')
# ------------------------------------------------------------
if ss.stage == "review":
    st.subheader("Step 2 — Review candidate gaps (accept / edit)")

    raw_items: List[Dict[str, Any]] = []
    for a in ss.per_paper:
        pid = a.get("paper_id") or "unknown"
        for item in a.get("reported_limitations", []) or []:
            if isinstance(item, dict):
                txt = (item.get("description") or "").strip()
                if txt:
                    raw_items.append({"paper_id": pid, "text": txt, "source": "Reported", "severity": 3,
                                      "section": item.get("section", ""), "evidence_quote": item.get("evidence_quote", ""),
                                      "evidence_page": item.get("evidence_page", None), "tags": []})
            elif isinstance(item, (str, int, float)):
                t = str(item).strip()
                if t:
                    raw_items.append({"paper_id": pid, "text": t, "source": "Reported", "severity": 3,
                                      "section": "", "evidence_quote": "", "evidence_page": None, "tags": []})
        for g in a.get("inferred_gaps", []) or []:
            if isinstance(g, dict):
                desc = (g.get("description") or "").strip()
                if desc:
                    sev = parse_severity(g.get("severity", 3))
                    raw_items.append({"paper_id": pid, "text": desc, "source": "Inferred",
                                      "severity": min(max(sev, 1), 5), "section": "",
                                      "evidence_quote": g.get("evidence_quote", ""), "evidence_page": g.get("evidence_page", None),
                                      "tags": []})
            elif isinstance(g, (str, int, float)):
                t = str(g).strip()
                if t:
                    raw_items.append({"paper_id": pid, "text": t, "source": "Inferred", "severity": 3,
                                      "section": "", "evidence_quote": "", "evidence_page": None, "tags": []})

    if not raw_items:
        st.warning("No gap statements found.")
        st.stop()

    if tagging_on:
        try:
            client = build_client(api_key) if api_key else None
            if client:
                with st.spinner("Tagging gaps…"):
                    gap_texts = [x["text"] for x in raw_items]
                    tag_lists = llm_tag_gaps(client, model, gap_texts)
                    for x, tags in zip(raw_items, tag_lists):
                        x["tags"] = [t for t in tags if t in CHECKLIST][:4]
        except Exception:
            pass

    df = pd.DataFrame([{
        "Accept": True,
        "Text": item["text"],
        "Source": item["source"],
        "Severity": int(item.get("severity", 3)),
        "Section": item.get("section", ""),
        "EvidenceQuote": item.get("evidence_quote", ""),
        "EvidencePage": item.get("evidence_page", ""),
        "Tags": ", ".join(item.get("tags", [])),
        "PaperID": item["paper_id"]
    } for item in raw_items])

    st.caption("Uncheck to exclude; edit text/evidence freely before clustering.")
    edited = st.data_editor(
        df, height=360, width="stretch",
        column_config={
            "Accept": st.column_config.CheckboxColumn(),
            "Text": st.column_config.TextColumn(width="large"),
            "Severity": st.column_config.NumberColumn(min_value=1, max_value=5, step=1),
        },
        key="editor_grid"
    )

    st.download_button(
        "⬇️ Download CSV of current (edited) gaps",
        data=edited.to_csv(index=False).encode("utf-8"),
        file_name="accepted_gaps_edited.csv",
        mime="text/csv",
        width="stretch"
    )

    col_c1, col_c2 = st.columns([1, 1])
    proceed = col_c1.button("Cluster accepted gaps", type="primary", width="stretch")
    fast_two_paras = col_c2.button("➡️ Skip review and generate Two-Paragraph report", type="secondary", width="stretch")

    def _collect_accepted_from_editor() -> List[Dict[str, Any]]:
        accepted_list = []
        for row, base in zip(edited.to_dict("records"), raw_items):
            if not row.get("Accept", False):
                continue
            item = base.copy()
            item["text"] = (row.get("Text") or "").strip()
            item["severity"] = int(row.get("Severity") or 3)
            item["section"] = row.get("Section", "")
            item["evidence_quote"] = row.get("EvidenceQuote", "")
            item["evidence_page"] = row.get("EvidencePage", "")
            item["tags"] = [t.strip() for t in (row.get("Tags", "") or "").split(",") if t.strip()]
            accepted_list.append(item)
        return accepted_list

    if proceed or fast_two_paras:
        accepted = _collect_accepted_from_editor()
        if not accepted:
            st.info("No rows selected; using all items for clustering.")
            accepted = []
            for base in raw_items:
                item = base.copy()
                item["severity"] = int(item.get("severity", 3))
                accepted.append(item)
        ss.accepted = accepted
        ss.stage = "cluster"
        _safe_rerun()
    else:
        st.stop()

# ------------------------------------------------------------
# Step 3 — Clustering
# ------------------------------------------------------------
if ss.stage == "cluster":
    st.subheader("Step 3 — Clustering shared gaps")

    if not ss.accepted:
        st.warning("No accepted gaps. Go back to review.")
        st.stop()

    client = build_client(api_key)
    texts = [g["text"] for g in ss.accepted]

    st.info(f"Embedding {len(texts)} items with {embedding_model} …")
    vecs = embed(client, embedding_model, texts)

    st.info("Clustering …")
    clusters_members = hdbscan_cluster(vecs) if (rigor and HDBSCAN_AVAILABLE and sk_normalize is not None) \
                      else greedy_cluster(vecs, threshold=sim_threshold)

    stats: List[Dict[str, Any]] = [cluster_stats_from_members(mem, ss.accepted) for mem in clusters_members]
    order = rank_clusters(stats, total_papers=len(ss.per_paper), rigorous=rigor)
    top_idx = order[0] if order else 0

    payload: Dict[str, Any] = {"total_papers": len(ss.per_paper), "clusters": []}
    appendix_items: List[Dict[str, Any]] = []
    for ci in order:
        c = stats[ci]
        payload["clusters"].append({
            "index": ci, "label": f"Cluster {ci}",
            "paper_coverage": c["paper_coverage"], "count": c["count"],
            "avg_severity": round(c["avg_severity"], 2), "median_severity": round(c["median_severity"], 2),
            "reported_share": round(c["reported_share"], 3), "inferred_share": round(c["inferred_share"], 3),
            "evidence_strength": round(c["evidence_strength"], 3),
            "example_statements": c["samples"], "is_recommended": (ci == top_idx)
        })
        for idx in c["members_idx"]:
            it = ss.accepted[idx]
            if it.get("evidence_quote"):
                appendix_items.append({
                    "paper_id": it.get("paper_id", "unknown"),
                    "evidence_quote": it.get("evidence_quote", ""),
                    "evidence_page": it.get("evidence_page", "")
                })
            if len(appendix_items) >= 40:
                break

    st.success(f"Formed {len(stats)} clusters. Recommended: **Cluster {top_idx}**")
    with st.expander("Preview: top cluster examples"):
        st.write(stats[order[0]]["samples"] if order else [])

    ss.stats = stats
    ss.order = order
    ss.payload = payload
    ss.appendix = appendix_items
    ss.stage = "report"
    _safe_rerun()

# ------------------------------------------------------------
# Step 4 — Report
# ------------------------------------------------------------
if ss.stage == "report":
    st.subheader("Step 4 — Generate report")

    style = st.radio(
        "Choose your main report style",
        ["Two-paragraph (recommended)", "Narrative", "Executive Brief (short)", "Plain-Language (very simple words)"],
        index=0,
        horizontal=True
    )

    client = build_client(api_key)

    try:
        if style == "Two-paragraph (recommended)":
            p1, p2 = llm_two_paragraph_report(client, model, ss.payload)
            main_text = f"{p1}\n\n{p2}"
            main_title = "Two-Paragraph Summary — Gaps and Next Steps"
            main_fname = "two_paragraph_summary.docx"
            paragraphs = [p1, p2]
        elif style == "Executive Brief (short)":
            main_text = llm_executive_brief(client, model, ss.payload)
            main_title = "Executive Brief — Research Gaps"
            main_fname = "executive_brief.docx"
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", main_text) if p.strip()]
        elif style == "Plain-Language (very simple words)":
            main_text = llm_plain_language(client, model, ss.payload)
            main_title = "Plain-Language Summary — What’s Missing and What To Do"
            main_fname = "plain_language_summary.docx"
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", main_text) if p.strip()]
        else:
            main_text = llm_narrative(client, model, ss.payload)
            main_title = "Research Gaps — Narrative Summary"
            main_fname = "final_report_narrative.docx"
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", main_text) if p.strip()]
    except Exception as e:
        st.error(f"Report generation failed: {e}")
        st.stop()

    st.text_area("Report preview", value="\n\n".join(paragraphs), height=260, width="stretch")

    if Document is not None:
        data_main = docx_from_paragraphs_with_appendix(main_title, paragraphs, ss.appendix, logo_bytes, accent)
        main_ext = ".docx"; main_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        data_main = rtf_from_paragraphs(main_title, paragraphs)
        main_ext = ".rtf"; main_mime = "application/rtf"

    data_txt = txt_from_paragraphs(main_title, paragraphs)

    st.download_button(
        f"⬇️ Download {main_fname if main_ext=='.docx' else 'report.rtf'}",
        data=data_main,
        file_name=(main_fname if main_ext == ".docx" else "report.rtf"),
        mime=main_mime,
        width="stretch"
    )
    st.download_button(
        "⬇️ Download TXT (always available)",
        data=data_txt,
        file_name="report.txt",
        mime="text/plain",
        width="stretch"
    )

    # CSV of accepted gaps (final snapshot)
    final_df = pd.DataFrame([{
        "PaperID": it["paper_id"],
        "Text": it["text"],
        "Source": it["source"],
        "Severity": it["severity"],
        "Section": it.get("section", ""),
        "EvidenceQuote": it.get("evidence_quote", ""),
        "EvidencePage": it.get("evidence_page", ""),
        "Tags": ", ".join(it.get("tags", []))
    } for it in ss.accepted])

    st.download_button(
        "⬇️ Download CSV of accepted gaps (final)",
        data=final_df.to_csv(index=False).encode("utf-8"),
        file_name="accepted_gaps_final.csv",
        mime="text/csv",
        width="stretch"
    )
